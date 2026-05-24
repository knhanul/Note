"""Folder import service.

Walks a source directory tree on disk, mirrors the folder hierarchy into the
current library's note DB, and creates a note record for each supported
document file (md/markdown/txt/html/htm/docx). Relative images referenced from
.md files are inlined as base64 data URLs so the imported notes render
self-contained inside the editor.
"""
from __future__ import annotations

import base64
import html as html_mod
import os
import re
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from services.folder_service import FolderService
from services.markdown_asset_resolver import extract_markdown_assets
from services.note_service import NoteService


_MD_IMG_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_HTML_BR_PATTERN = re.compile(r"<br\s*/?>", re.IGNORECASE)
_HTML_BLOCK_CLOSE_PATTERN = re.compile(
    r"</(p|div|li|ul|ol|h[1-6]|section|article|blockquote|pre)>",
    re.IGNORECASE,
)
_HTML_TAG_PATTERN = re.compile(r"<[^>]+>")

_EXT_TO_MIME = {
    "png": "png",
    "jpg": "jpeg",
    "jpeg": "jpeg",
    "gif": "gif",
    "webp": "webp",
    "bmp": "bmp",
    "svg": "svg+xml",
}


class FolderImportService:
    """Imports a directory tree of documents into the current library."""

    SUPPORTED_EXTS = {".md", ".markdown", ".txt", ".html", ".htm", ".docx", ".hwp", ".hwpx"}
    IMPORT_MODE_FAST_TEXT = "fast_text"
    IMPORT_MODE_STRUCTURED = "structured"
    IMPORT_MODE_AUTO = "auto"
    DEFAULT_IMPORT_MODE = IMPORT_MODE_AUTO

    def __init__(
        self,
        folder_service: FolderService,
        note_service: NoteService,
    ) -> None:
        self._folders = folder_service
        self._notes = note_service

    def import_directory(
        self,
        src_dir: str,
        parent_folder_id: Optional[str] = None,
        folder_color: str = "#3B82F6",
        include_subfolders: bool = True,
        import_mode: str = DEFAULT_IMPORT_MODE,
        progress_callback=None,
    ) -> Dict[str, Any]:
        """Import documents from a directory tree.

        import_mode:
        - fast_text: legacy gethwp-based fast text extraction
        - structured: structured HWPX importer path only
        - auto: structured first, fallback to fast_text
        """
        if not src_dir:
            raise ValueError("가져올 폴더가 지정되지 않았습니다.")
        mode = self._normalize_import_mode(import_mode)
        if mode != (import_mode or self.DEFAULT_IMPORT_MODE).strip().lower():
            print(f"[FolderImport] Unknown import_mode='{import_mode}', fallback to auto")

        print(f"[FolderImport] import_mode={mode}")

        src = Path(src_dir)
        if not src.exists() or not src.is_dir():
            raise ValueError(f"폴더를 찾을 수 없습니다: {src_dir}")

        # When not including subfolders, import files directly to parent without creating root folder
        if not include_subfolders:
            print(f"[FolderImport] Importing files only (no folder creation)")
            imported_notes: List[str] = []
            failures: List[Dict[str, str]] = []

            # If no parent folder, create a default folder
            target_folder_id = parent_folder_id
            if not target_folder_id:
                default_name = "가져온 문서"
                default_color = "#3B82F6"
                target_folder_id = self._create_folder(default_name, None, default_color)
                if not target_folder_id:
                    raise RuntimeError("기본 폴더 생성에 실패했습니다.")
                print(f"[FolderImport] Created default folder: '{default_name}' -> {target_folder_id}")
                created_folders = [target_folder_id]
            else:
                created_folders = []

            # Process only files in the root directory
            files = sorted([f for f in src.iterdir() if f.is_file()])
            total_files = len(files)
            print(f"[FolderImport] Processing {total_files} files in root directory")
            processed = 0

            for fpath in files:
                processed += 1
                self._process_file(
                    fpath, target_folder_id, processed, total_files,
                    mode, progress_callback, imported_notes, failures,
                )

            print(f"[FolderImport] Import complete: {len(imported_notes)} notes, 0 folders, {len(failures)} failures")
            if failures:
                print(f"[FolderImport] Failures:")
                for f in failures:
                    print(f"[FolderImport]   - {f['path']}: {f['error']}")

            return {
                "rootFolderId": target_folder_id or "",
                "rootLabel": "",
                "noteCount": len(imported_notes),
                "folderCount": 0,
                "failedCount": len(failures),
                "failures": failures,
            }

        # Original behavior with folder structure
        root_label = src.name or "가져온 폴더"
        print(f"[FolderImport] Creating root folder: '{root_label}' (include_subfolders={include_subfolders})")
        root_folder_id = self._create_folder(root_label, parent_folder_id, folder_color)
        if not root_folder_id:
            raise RuntimeError("최상위 폴더 생성에 실패했습니다.")
        print(f"[FolderImport] Root folder created: {root_folder_id}")

        path_to_folder: Dict[Path, str] = {src.resolve(): root_folder_id}
        imported_notes: List[str] = []
        created_folders: List[str] = [root_folder_id]
        failures: List[Dict[str, str]] = []

        # ── Pre-scan: directories that contain (or lead to) supported files ──
        dirs_with_files: set = set()
        for _current_dir, _sub_dirs, files in os.walk(src):
            _current_path = Path(_current_dir).resolve()
            if any((_current_path / f).suffix.lower() in self.SUPPORTED_EXTS for f in files):
                p = _current_path
                while p and p != src.parent.resolve():
                    dirs_with_files.add(p)
                    p = p.parent.resolve()

        # Count total files for progress
        total_files = 0
        for _current_dir, _sub_dirs, files in os.walk(src):
            for fname in files:
                fpath = Path(_current_dir) / fname
                if fpath.suffix.lower() in self.SUPPORTED_EXTS:
                    total_files += 1
        processed = 0

        for current_dir, sub_dirs, files in os.walk(src):
            current_path = Path(current_dir).resolve()
            current_folder_id = path_to_folder.get(current_path)
            if current_folder_id is None:
                sub_dirs[:] = []
                continue

            sub_dirs.sort()
            files.sort()

            if not include_subfolders and current_path != src.resolve():
                sub_dirs[:] = []
                continue

            print(f"[FolderImport] Processing directory: {current_path} ({len(sub_dirs)} subdirs, {len(files)} files)")

            # Only create sub-folders that are on a path to supported files
            kept_subs: List[str] = []
            for sub in sub_dirs:
                sub_path = (current_path / sub).resolve()
                if sub_path not in dirs_with_files:
                    print(f"[FolderImport] Skipping empty folder: '{sub}'")
                    continue
                new_id = self._create_folder(sub, current_folder_id, folder_color)
                if new_id:
                    path_to_folder[sub_path] = new_id
                    created_folders.append(new_id)
                    kept_subs.append(sub)
                    print(f"[FolderImport] Created folder: '{sub}' -> {new_id}")
                else:
                    failures.append({"path": str(sub_path), "error": "폴더 생성 실패"})
                    print(f"[FolderImport] Failed to create folder: '{sub}'")
            sub_dirs[:] = kept_subs

            # Import files
            for fname in files:
                fpath = current_path / fname
                processed += 1
                self._process_file(
                    fpath, current_folder_id, processed, total_files,
                    mode, progress_callback, imported_notes, failures,
                )

        print(f"[FolderImport] Import complete: {len(imported_notes)} notes, {len(created_folders)} folders, {len(failures)} failures")
        if failures:
            print(f"[FolderImport] Failures:")
            for f in failures:
                print(f"[FolderImport]   - {f['path']}: {f['error']}")

        return {
            "rootFolderId": root_folder_id,
            "rootLabel": root_label,
            "noteCount": len(imported_notes),
            "folderCount": len(created_folders),
            "failedCount": len(failures),
            "failures": failures,
        }

    # ── helpers ────────────────────────────────────────────────────────────
    def _process_file(
        self,
        fpath: Path,
        folder_id: str,
        processed: int,
        total_files: int,
        mode: str,
        progress_callback,
        imported_notes: List[str],
        failures: List[Dict[str, str]],
    ) -> None:
        """Process a single file: read, convert, and create a note record."""
        ext = fpath.suffix.lower()
        if ext not in self.SUPPORTED_EXTS:
            print(f"[FolderImport] Skipping file (unsupported ext): {fpath.name} ({ext})")
            if progress_callback:
                progress_callback(processed, total_files, f"스킵: {fpath.name}")
            return

        print(f"[FolderImport] Processing file: {fpath.name} ({ext})")
        if progress_callback:
            progress_callback(processed, total_files, f"읽는 중: {fpath.name}")
        try:
            title, markdown, tags = self._read_note(fpath, import_mode=mode)
        except Exception as exc:  # noqa: BLE001
            failures.append({"path": str(fpath), "error": str(exc)})
            print(f"[FolderImport] Failed to read file {fpath.name}: {exc}")
            if progress_callback:
                progress_callback(processed, total_files, f"오류: {fpath.name}")
            return

        note_id = uuid.uuid4().hex[:8]
        if progress_callback:
            progress_callback(processed, total_files, f"저장 중: {fpath.name}")
        if self._notes.create(
            note_id=note_id,
            folder_id=folder_id,
            title=title or fpath.stem or "무제",
            content=markdown or "",
            content_json="",
            tags=tags,
        ):
            imported_notes.append(note_id)
            print(f"[FolderImport] Imported note: '{title or fpath.stem}' -> {note_id}")
        else:
            failures.append({"path": str(fpath), "error": "노트 생성 실패"})
            print(f"[FolderImport] Failed to create note: {fpath.name}")

    def _create_folder(
        self, name: str, parent_id: Optional[str], color: str
    ) -> Optional[str]:
        clean = (name or "").strip() or "무제 폴더"
        folder_id = uuid.uuid4().hex[:8]
        ok = self._folders.create(folder_id, clean, color, parent_id)
        return folder_id if ok else None

    def _read_note(self, fpath: Path, import_mode: str = DEFAULT_IMPORT_MODE) -> Tuple[str, str, List[str]]:
        from packages.import_export.markdown_import_service import load_markdown_document_from_text
        from packages.import_export.hwp_import_service import convert_hwp_to_markdown_text
        from packages.import_export.hwpx_import_service import convert_hwpx_to_markdown_text

        ext = fpath.suffix.lower()
        title = fpath.stem
        filename_line = f"# {fpath.stem}\n\n"
        tags = []
        if ext in (".md", ".markdown"):
            text = self._read_text(fpath)
            doc, asset_warnings = load_markdown_document_from_text(text, source_path=str(fpath))
            if doc.warnings:
                for w in doc.warnings:
                    print(f"[FolderImport] Warning: {w}")
            for w in asset_warnings:
                print(f"[FolderImport] Warning: {w}")
            title = doc.metadata.title if doc.metadata.title else fpath.stem
            filename_line = "" if doc.metadata.title else f"# {fpath.stem}\n\n"
            tags = doc.metadata.tags if doc.metadata.tags else []
            return title, filename_line + self._inline_md_images(doc.body_markdown, fpath.parent), tags
        if ext == ".txt":
            return title, filename_line + self._read_text(fpath), tags
        if ext in (".html", ".htm"):
            return title, filename_line + self._html_to_markdown(self._read_text(fpath)), tags
        if ext == ".docx":
            return title, filename_line + self._docx_to_markdown(fpath), tags
        if ext == ".hwp":
            markdown, warnings = convert_hwp_to_markdown_text(str(fpath))
            for w in warnings:
                print(f"[FolderImport] HWP Warning: {w}")
            return title, filename_line + markdown, tags
        if ext == ".hwpx":
            markdown, warnings = convert_hwpx_to_markdown_text(str(fpath))
            for w in warnings:
                print(f"[FolderImport] HWPX Warning: {w}")
            return title, filename_line + markdown, tags
        return title, filename_line, tags

    @staticmethod
    def _read_text(fpath: Path) -> str:
        for enc in ("utf-8", "utf-8-sig", "cp949", "euc-kr", "latin-1"):
            try:
                return fpath.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        return fpath.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _parse_markdown_metadata(text: str) -> Tuple[Optional[Dict[str, Any]], str]:
        """Parse metadata from markdown frontmatter and return (metadata, body)."""
        if not text:
            return None, text

        lines = text.split("\n")
        if not lines:
            return None, text

        # Check for YAML frontmatter (---)
        if lines[0].strip() == "---":
            end_index = -1
            for i in range(1, len(lines)):
                if lines[i].strip() == "---":
                    end_index = i
                    break

            if end_index > 0:
                try:
                    import yaml
                    metadata_text = "\n".join(lines[1:end_index])
                    metadata = yaml.safe_load(metadata_text) or {}
                    body = "\n".join(lines[end_index + 1:])
                    return metadata, body
                except Exception:
                    # YAML parsing failed, treat as regular markdown
                    pass

        # Check for TOML frontmatter (+++)
        if lines[0].strip() == "+++":
            end_index = -1
            for i in range(1, len(lines)):
                if lines[i].strip() == "+++":
                    end_index = i
                    break

            if end_index > 0:
                try:
                    import toml
                    metadata_text = "\n".join(lines[1:end_index])
                    metadata = toml.loads(metadata_text)
                    body = "\n".join(lines[end_index + 1:])
                    return metadata, body
                except Exception:
                    # TOML parsing failed, treat as regular markdown
                    pass

        # No metadata found
        return None, text

    @staticmethod
    def _inline_md_images(markdown: str, base_dir: Path) -> str:
        def _replace(match: re.Match) -> str:
            alt = match.group(1)
            src = (match.group(2) or "").strip()
            if not src or src.startswith(("data:", "http://", "https://")):
                return match.group(0)

            # Strip optional title segment: ![alt](url "title")
            url_only = src.split(" ", 1)[0]

            try:
                img_path = (base_dir / url_only).resolve()
                if not img_path.exists() or not img_path.is_file():
                    return match.group(0)
                ext = img_path.suffix.lower().lstrip(".")
                mime = _EXT_TO_MIME.get(ext)
                if not mime:
                    return match.group(0)
                data = base64.b64encode(img_path.read_bytes()).decode("ascii")
                return f"![{alt}](data:image/{mime};base64,{data})"
            except Exception:
                return match.group(0)

        return _MD_IMG_PATTERN.sub(_replace, markdown or "")

    @staticmethod
    def _html_to_markdown(raw_html: str) -> str:
        text = _HTML_BR_PATTERN.sub("\n", raw_html or "")
        text = _HTML_BLOCK_CLOSE_PATTERN.sub("\n", text)
        text = _HTML_TAG_PATTERN.sub("", text)
        text = html_mod.unescape(text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _docx_to_markdown(fpath: Path) -> str:
        try:
            from docx import Document
        except Exception:  # noqa: BLE001
            return ""

        doc = Document(str(fpath))
        blocks: List[str] = []
        for para in doc.paragraphs:
            text = (para.text or "").rstrip()
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                m = re.search(r"\d+", style)
                level = max(1, min(6, int(m.group()) if m else 1))
                blocks.append(("#" * level) + " " + text)
            elif "list" in style and text.strip():
                blocks.append("- " + text)
            else:
                blocks.append(text)

        for table in doc.tables:
            rows: List[List[str]] = []
            for row in table.rows:
                rows.append(
                    [cell.text.replace("\n", " ").strip() for cell in row.cells]
                )
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            rows = [r + [""] * (cols - len(r)) for r in rows]
            md_lines = ["| " + " | ".join(rows[0]) + " |"]
            md_lines.append("| " + " | ".join(["---"] * cols) + " |")
            for r in rows[1:]:
                md_lines.append("| " + " | ".join(r) + " |")
            blocks.append("\n".join(md_lines))

        joined = "\n\n".join(b for b in blocks if b is not None)
        return re.sub(r"\n{3,}", "\n\n", joined).strip()

    @staticmethod
    def _hwp_to_markdown(fpath: Path, import_mode: str = DEFAULT_IMPORT_MODE) -> str:
        """Convert HWP/HWPX to markdown using mode-aware fallback chain.

        Developer flow:
        1) structured path (.hwpx direct, .hwp via COM->HWPX)
        2) fast_text path (legacy gethwp)
        """
        ext = fpath.suffix.lower()
        mode = FolderImportService._normalize_import_mode(import_mode)
        print(f"[FolderImport] Reading HWP file: {fpath} (mode={mode})")

        use_structured = mode in {
            FolderImportService.IMPORT_MODE_STRUCTURED,
            FolderImportService.IMPORT_MODE_AUTO,
        }
        use_fast_text = mode in {
            FolderImportService.IMPORT_MODE_FAST_TEXT,
            FolderImportService.IMPORT_MODE_AUTO,
        }

        if use_structured:
            structured_markdown = FolderImportService._try_structured_import(fpath, ext)
            if structured_markdown:
                return structured_markdown

        if not use_fast_text:
            print("[FolderImport] fast_text path disabled by import_mode")
            return ""

        return FolderImportService._read_via_gethwp(fpath, ext)

    @staticmethod
    def _normalize_import_mode(import_mode: str) -> str:
        mode = (import_mode or FolderImportService.DEFAULT_IMPORT_MODE).strip().lower()
        if mode in {
            FolderImportService.IMPORT_MODE_FAST_TEXT,
            FolderImportService.IMPORT_MODE_STRUCTURED,
            FolderImportService.IMPORT_MODE_AUTO,
        }:
            return mode
        return FolderImportService.DEFAULT_IMPORT_MODE

    @staticmethod
    def _try_structured_import(fpath: Path, ext: str) -> str:
        if ext == ".hwpx":
            return FolderImportService._parse_hwpx_with_importer(fpath)
        if ext == ".hwp":
            return FolderImportService._parse_hwp_via_com_then_importer(fpath)
        return ""

    @staticmethod
    def _parse_hwpx_with_importer(hwpx_path: Path) -> str:
        try:
            from services.hwpx_importer import hwpx_to_markdown

            parsed = (hwpx_to_markdown(str(hwpx_path)) or "").strip()
            if parsed:
                print("[FolderImport] HWPX importer success")
                return parsed
            print("[FolderImport] HWPX importer failed, fallback to gethwp")
            return ""
        except Exception as exc:  # noqa: BLE001
            print(f"[FolderImport] HWPX importer failed, fallback to gethwp: {exc}")
            return ""

    @staticmethod
    def _parse_hwp_via_com_then_importer(hwp_path: Path) -> str:
        converted_hwpx: Path | None = None
        try:
            from services.hwp_converter import convert_hwp_to_hwpx_via_com

            converted = convert_hwp_to_hwpx_via_com(str(hwp_path))
            if not converted:
                print("[FolderImport] HWP COM conversion failed, fallback to gethwp")
                return ""

            converted_hwpx = Path(converted)
            return FolderImportService._parse_hwpx_with_importer(converted_hwpx)
        except Exception as exc:  # noqa: BLE001
            print(f"[FolderImport] HWP COM/importer failed, fallback to gethwp: {exc}")
            return ""
        finally:
            FolderImportService._cleanup_temp_hwpx(converted_hwpx)

    @staticmethod
    def _cleanup_temp_hwpx(hwpx_path: Path | None) -> None:
        if hwpx_path is None:
            return
        try:
            parent_dir = hwpx_path.parent
            if hwpx_path.exists():
                hwpx_path.unlink()
            if parent_dir.exists() and parent_dir.name.startswith("hwp_import_"):
                parent_dir.rmdir()
        except Exception:
            pass

    @staticmethod
    def _read_via_gethwp(fpath: Path, ext: str) -> str:

        try:
            import gethwp
        except Exception as exc:  # noqa: BLE001
            print(f"[FolderImport] gethwp unavailable: {exc}")
            return ""

        try:
            if ext == ".hwp":
                text = gethwp.read_hwp(str(fpath))
            elif ext == ".hwpx":
                text = gethwp.read_hwpx(str(fpath))
            else:
                return ""

            if not text:
                print(f"[FolderImport] HWP file returned empty text")
                return ""

            lines = text.split("\n")
            blocks = []
            for line in lines:
                line = line.strip()
                if line:
                    blocks.append(line)

            joined = "\n\n".join(b for b in blocks if b)
            print(f"[FolderImport] HWP conversion complete: {len(blocks)} paragraphs, {len(joined)} chars")
            return re.sub(r"\n{3,}", "\n\n", joined).strip()
        except Exception as exc:
            print(f"[FolderImport] HWP conversion failed: {exc}")
            return ""
