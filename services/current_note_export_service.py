"""Single-note export service for editor-driven high-quality exports."""
from __future__ import annotations

import base64
import html
import re
import subprocess
import sys
import tempfile
from io import BytesIO
from pathlib import Path
from typing import List, Tuple


_DATA_URL_PATTERN = re.compile(
    r"data:(image/[a-zA-Z0-9.+-]+);base64,([A-Za-z0-9+/=\r\n]+)"
)
_MD_IMG_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_INVALID_FS_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
_TABLE_SEP_PATTERN = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
_HTML_TABLE_PATTERN = re.compile(r"<table\b[^>]*>.*?</table>", re.IGNORECASE | re.DOTALL)
_HTML_TR_PATTERN = re.compile(r"<tr\b[^>]*>(.*?)</tr>", re.IGNORECASE | re.DOTALL)
_HTML_CELL_PATTERN = re.compile(r"<t[hd]\b[^>]*>(.*?)</t[hd]>", re.IGNORECASE | re.DOTALL)
_HTML_BR_PATTERN = re.compile(r"<br\s*/?>", re.IGNORECASE)
_HTML_BLOCK_CLOSE_PATTERN = re.compile(
    r"</(p|div|li|ul|ol|h[1-6]|section|article|blockquote|pre)>",
    re.IGNORECASE,
)
_HTML_TAG_PATTERN = re.compile(r"<[^>]+>")


class CurrentNoteExportService:
    """Exports a single note (currently opened in editor) to multiple formats."""

    SUPPORTED_FORMATS = ("md", "txt", "pdf", "hwpx", "docx")
    HWPX_FALLBACK_FORMAT = "docx"

    @staticmethod
    def safe_filename(name: str) -> str:
        cleaned = _INVALID_FS_CHARS.sub("_", (name or "").strip()).strip(" .")
        return (cleaned or "무제")[:120]

    def export(
        self,
        title: str,
        markdown: str,
        content_json: str,
        fmt: str,
        out_dir: str,
    ) -> str:
        fmt = (fmt or "").lower().strip()
        if fmt not in self.SUPPORTED_FORMATS:
            raise ValueError(f"지원하지 않는 포맷: {fmt}")

        target_dir = Path(out_dir)
        target_dir.mkdir(parents=True, exist_ok=True)

        base_name = self.safe_filename(title or "무제")
        normalized_markdown = self._normalize_markdown_for_text_exports(markdown or "")
        if fmt == "md":
            return self._export_markdown(target_dir, base_name, normalized_markdown)
        if fmt == "txt":
            return self._export_txt(target_dir, base_name, title or "", normalized_markdown)
        if fmt == "docx":
            return self._export_docx(target_dir, base_name, title or "", markdown or "")
        if fmt == "hwpx":
            return self._export_hwpx_via_md2hwpx(target_dir, base_name, title or "", markdown or "")
        # pdf is handled in QML WebEngine print path for best WYSIWYG quality.
        raise ValueError("PDF export is handled by WebEngine path")

    def _export_markdown(self, target_dir: Path, base_name: str, markdown: str) -> str:
        counter = {"n": 0}

        def _replace(match: re.Match) -> str:
            alt = match.group(1)
            src = match.group(2).strip()
            if not src.startswith("data:"):
                return match.group(0)

            data_match = _DATA_URL_PATTERN.match(src)
            if not data_match:
                return match.group(0)

            mime = data_match.group(1)
            b64 = data_match.group(2).replace("\n", "").replace("\r", "")
            ext = _mime_to_ext(mime)

            counter["n"] += 1
            img_name = f"{base_name}_img_{counter['n']:03d}.{ext}"
            img_path = target_dir / img_name
            try:
                img_path.write_bytes(base64.b64decode(b64))
            except Exception:
                return match.group(0)
            return f"![{alt}]({img_name})"

        rewritten = _MD_IMG_PATTERN.sub(_replace, markdown)
        out_path = target_dir / f"{base_name}.md"
        out_path.write_text(rewritten, encoding="utf-8")
        return str(out_path)

    def _export_txt(self, target_dir: Path, base_name: str, title: str, markdown: str) -> str:
        txt = self._markdown_to_plain_text(markdown)
        body = (title + "\n\n" if title else "") + txt
        out_path = target_dir / f"{base_name}.txt"
        out_path.write_text(body, encoding="utf-8")
        return str(out_path)

    def _export_docx(self, target_dir: Path, base_name: str, title: str, markdown: str) -> str:
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError("python-docx가 설치되어 있지 않습니다.") from exc

        doc = Document()
        if title:
            doc.add_heading(title, level=1)

        lines = (markdown or "").splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            if _is_table_header(lines, i):
                table_lines, consumed = _collect_table_lines(lines, i)
                rows = [_split_table_row(row) for row in table_lines if row.strip()]
                if rows:
                    cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=cols)
                    for r_idx, row in enumerate(rows):
                        for c_idx in range(cols):
                            table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""
                i += consumed
                continue

            img_match = _MD_IMG_PATTERN.fullmatch(line.strip())
            if img_match:
                src = img_match.group(2).strip()
                data_match = _DATA_URL_PATTERN.match(src)
                if data_match:
                    b64 = data_match.group(2).replace("\n", "").replace("\r", "")
                    try:
                        doc.add_picture(BytesIO(base64.b64decode(b64)), width=Inches(5.8))
                    except Exception:
                        doc.add_paragraph(f"[이미지:{img_match.group(1) or 'image'}]")
                else:
                    doc.add_paragraph(f"[이미지:{img_match.group(1) or 'image'}]")
                i += 1
                continue

            heading_match = re.match(r"^(#{1,3})\s+(.*)$", line)
            if heading_match:
                level = min(len(heading_match.group(1)), 3)
                doc.add_heading(_strip_inline_md(heading_match.group(2)), level=level)
                i += 1
                continue

            bullet_match = re.match(r"^\s*([-*+]|\d+\.)\s+(.*)$", line)
            if bullet_match:
                style = "List Number" if bullet_match.group(1).endswith(".") else "List Bullet"
                doc.add_paragraph(_strip_inline_md(bullet_match.group(2)), style=style)
                i += 1
                continue

            if line.strip():
                doc.add_paragraph(_strip_inline_md(line))
            else:
                doc.add_paragraph("")
            i += 1

        out_path = target_dir / f"{base_name}.docx"
        doc.save(str(out_path))
        return str(out_path)

    def _export_hwpx_via_docx(self, target_dir: Path, base_name: str, title: str, markdown: str) -> str:
        docx_path = Path(self._export_docx(target_dir, base_name, title, markdown))
        hwpx_path = target_dir / f"{base_name}.hwpx"
        self._convert_docx_to_hwpx(docx_path, hwpx_path)
        return str(hwpx_path)

    def _export_hwpx_via_md2hwpx(self, target_dir: Path, base_name: str, title: str, markdown: str) -> str:
        hwpx_path = target_dir / f"{base_name}.hwpx"
        source_md = (f"# {title}\n\n" if title else "") + self._normalize_markdown_for_text_exports(markdown or "")

        with tempfile.TemporaryDirectory(prefix="hwpx_export_") as temp_dir:
            temp_path = Path(temp_dir)
            temp_md_path = self._prepare_hwpx_markdown_assets(temp_path, base_name, source_md)
            commands = [
                [
                    "md2hwpx",
                    str(temp_md_path.name),
                    "-o",
                    str(hwpx_path),
                ],
                [
                    sys.executable,
                    "-m",
                    "md2hwpx",
                    str(temp_md_path.name),
                    "-o",
                    str(hwpx_path),
                ],
            ]
            errors: list[str] = []
            success = False

            for command in commands:
                try:
                    completed = subprocess.run(
                        command,
                        check=False,
                        capture_output=True,
                        text=True,
                        cwd=str(temp_path),
                    )
                except FileNotFoundError as exc:
                    errors.append(f"{' '.join(command)} -> 실행 파일 없음: {exc}")
                    continue

                if completed.returncode == 0:
                    success = True
                    break

                details = (completed.stderr or completed.stdout or "").strip()
                errors.append(f"{' '.join(command)} -> exit={completed.returncode}: {details}")

            if not success:
                raise RuntimeError(
                    "md2hwpx 변환에 실패했습니다. "
                    "`pip install md2hwpx` 설치 및 실행 환경을 확인해 주세요. "
                    f"상세: {' | '.join(errors)}"
                )

        if not hwpx_path.exists():
            raise RuntimeError("md2hwpx 변환이 완료되었지만 HWPX 파일이 생성되지 않았습니다.")

        return str(hwpx_path)

    def _prepare_hwpx_markdown_assets(self, target_dir: Path, base_name: str, markdown: str) -> Path:
        counter = {"n": 0}
        out_path = target_dir / f"{base_name}.md"

        def _replace(match: re.Match) -> str:
            alt = match.group(1)
            src = match.group(2).strip()
            if not src.startswith("data:"):
                return match.group(0)

            data_match = _DATA_URL_PATTERN.match(src)
            if not data_match:
                return match.group(0)

            mime = data_match.group(1)
            b64 = data_match.group(2).replace("\n", "").replace("\r", "")

            try:
                raw = base64.b64decode(b64)
            except Exception:
                return match.group(0)

            image_bytes = raw
            ext = "png"
            converted = False
            try:
                from PIL import Image

                with Image.open(BytesIO(raw)) as image:
                    converted_buffer = BytesIO()
                    image.convert("RGB").save(converted_buffer, format="PNG")
                image_bytes = converted_buffer.getvalue()
                converted = True
            except Exception:
                if _mime_to_ext(mime) in {"png", "jpg", "jpeg"}:
                    ext = _mime_to_ext(mime)
                else:
                    return f"[이미지 변환 실패:{alt or 'image'}]"

            counter["n"] += 1
            img_name = f"hwpx_img_{counter['n']:03d}.{'png' if converted else ext}"
            img_path = target_dir / img_name
            try:
                img_path.write_bytes(image_bytes)
            except Exception:
                return match.group(0)

            return f"![{alt}]({img_name})"

        rewritten = _MD_IMG_PATTERN.sub(_replace, markdown or "")
        out_path.write_text(rewritten, encoding="utf-8")
        return out_path

    @staticmethod
    def _convert_docx_to_hwpx(docx_path: Path, hwpx_path: Path) -> None:
        try:
            import pythoncom
            import win32com.client
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError("DOCX 후처리용 pywin32가 설치되어 있지 않습니다.") from exc

        abs_docx = str(docx_path.resolve())
        abs_hwpx = str(hwpx_path.resolve())

        pythoncom.CoInitialize()
        hwp = None
        errors: list[str] = []

        prog_ids = ["HWPFrame.HwpObject", "Hwp.HwpObject"]
        methods = [("DispatchEx", win32com.client.DispatchEx), ("Dispatch", win32com.client.Dispatch)]

        for prog_id in prog_ids:
            for method_name, dispatch in methods:
                try:
                    hwp = dispatch(prog_id)
                    break
                except Exception as exc:  # noqa: BLE001
                    errors.append(f"{method_name}('{prog_id}'): {exc}")
            else:
                continue
            break

        if hwp is None:
            pythoncom.CoUninitialize()
            raise RuntimeError(
                "한글 COM 객체를 생성하지 못했습니다. "
                f"시도한 ProgID: {prog_ids}. "
                f"상세 에러: {'; '.join(errors)}"
            )

        try:
            try:
                hwp.XHwpWindows.Item(0).Visible = False
            except Exception:
                pass
            try:
                hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
            except Exception:
                pass
            opened = False
            open_errors: list[str] = []
            for open_args in [("MSWORD", ""), ()]:
                try:
                    if open_args:
                        opened = bool(hwp.Open(abs_docx, *open_args))
                    else:
                        opened = bool(hwp.Open(abs_docx))
                    if opened:
                        break
                except Exception as exc:  # noqa: BLE001
                    open_errors.append(str(exc))
            if not opened:
                raise RuntimeError(
                    f"한글에서 DOCX 파일을 열지 못했습니다. "
                    f"상세 에러: {'; '.join(open_errors)}"
                )
            saved = bool(hwp.SaveAs(abs_hwpx, "HWPX", ""))
            if not saved or not hwpx_path.exists():
                raise RuntimeError("한글에서 HWPX 파일 저장에 실패했습니다.")
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(f"한글 자동화 후처리에 실패했습니다: {exc}") from exc
        finally:
            if hwp is not None:
                try:
                    hwp.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    @staticmethod
    def _markdown_to_plain_text(md: str) -> str:
        text = md or ""
        text = _MD_IMG_PATTERN.sub(lambda m: f"[이미지:{m.group(1) or 'image'}]", text)
        text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"^\s*[-*+]\s+", "• ", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+\.\s+", "• ", text, flags=re.MULTILINE)
        return text

    def _normalize_markdown_for_text_exports(self, markdown: str) -> str:
        text = markdown or ""

        def _table_to_md(match: re.Match) -> str:
            table_html = match.group(0)
            rows = []
            for tr_match in _HTML_TR_PATTERN.finditer(table_html):
                cell_values = []
                for cell_match in _HTML_CELL_PATTERN.finditer(tr_match.group(1)):
                    cell_text = _strip_html(cell_match.group(1))
                    cell_values.append(cell_text)
                if cell_values:
                    rows.append(cell_values)

            if not rows:
                return ""

            max_cols = max(len(r) for r in rows)
            rows = [r + [""] * (max_cols - len(r)) for r in rows]
            header = rows[0]
            sep = ["---"] * max_cols
            body = rows[1:]

            lines = [
                "| " + " | ".join(header) + " |",
                "| " + " | ".join(sep) + " |",
            ]
            for row in body:
                lines.append("| " + " | ".join(row) + " |")
            return "\n".join(lines)

        text = _HTML_TABLE_PATTERN.sub(_table_to_md, text)
        text = _HTML_BR_PATTERN.sub("\n", text)
        text = _HTML_BLOCK_CLOSE_PATTERN.sub("\n", text)
        text = _HTML_TAG_PATTERN.sub("", text)
        text = html.unescape(text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def _mime_to_ext(mime: str) -> str:
    mime = (mime or "").lower()
    if "png" in mime:
        return "png"
    if "jpeg" in mime or "jpg" in mime:
        return "jpg"
    if "gif" in mime:
        return "gif"
    if "webp" in mime:
        return "webp"
    if "bmp" in mime:
        return "bmp"
    if "svg" in mime:
        return "svg"
    return "png"


def _strip_inline_md(text: str) -> str:
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", text or "")
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    return cleaned


def _strip_html(raw: str) -> str:
    text = _HTML_BR_PATTERN.sub("\n", raw or "")
    text = _HTML_TAG_PATTERN.sub("", text)
    text = html.unescape(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _is_table_header(lines: List[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    line = lines[index].strip()
    if "|" not in line:
        return False
    return _TABLE_SEP_PATTERN.match(lines[index + 1].strip()) is not None


def _collect_table_lines(lines: List[str], start: int) -> Tuple[List[str], int]:
    collected: List[str] = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip() or "|" not in line:
            break
        collected.append(line)
        i += 1
    # skip markdown separator row if present in second row
    if len(collected) >= 2 and _TABLE_SEP_PATTERN.match(collected[1].strip()):
        collected.pop(1)
    return collected, i - start


def _split_table_row(line: str) -> List[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [_strip_inline_md(col.strip()) for col in raw.split("|")]
